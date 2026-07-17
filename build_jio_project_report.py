from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING

OUT = "Jio_FAQ_Chatbot_Project_Report.docx"

BLUE = "0B5CAD"; DARK = "163A5F"; MID = "2E74B5"; LIGHT = "EAF2FB"; GREY = "F2F4F7"; INK = "1F2937"; MUTED = "667085"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def borders(cell, color="D0D5DD"):
    tcPr = cell._tc.get_or_add_tcPr(); tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = tcBorders.find(qn(f'w:{edge}'))
        if e is None: e = OxmlElement(f'w:{edge}'); tcBorders.append(e)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),color)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def keep_with_next(p):
    pPr=p._p.get_or_add_pPr(); e=OxmlElement('w:keepNext'); pPr.append(e)

def add_page_field(p):
    r=p.add_run('Page '); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.addnext(fld)

def set_run(r, size=11, bold=False, color=INK, italic=False):
    r.font.name='Aptos'; r._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); r._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos')
    r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.color.rgb=RGBColor.from_string(color)

def add_para(doc, text='', style=None, before=0, after=6, align=None, size=None, bold=False, color=INK, italic=False):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.12
    if align is not None: p.alignment=align
    r=p.add_run(text); set_run(r, size or (11 if style is None else 11), bold, color, italic); return p

def add_heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.paragraph_format.keep_with_next=True; p.add_run(text); return p

def add_bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1; p.add_run(item)

def add_numbered(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1; p.add_run(item)

def add_table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,DARK); borders(c,DARK); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(h); set_run(r,9.2,True,'FFFFFF')
        if widths: c.width=Inches(widths[i])
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; borders(c); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: c.width=Inches(widths[i])
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
            r=p.add_run(str(val)); set_run(r,9.2)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def add_callout(doc, title, text):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,LIGHT); borders(c,'B8D4F0'); set_cell_margins(c,120,160,120,160)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); r=p.add_run(title); set_run(r,10.2,True,BLUE)
    p=c.add_paragraph(); p.paragraph_format.space_after=Pt(0); r=p.add_run(text); set_run(r,10)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def add_figure_placeholder(doc, number, description):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,'F8FAFC'); borders(c,'98A2B3'); set_cell_margins(c,420,120,420,120)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4); r=p.add_run('SCREENSHOT PLACEHOLDER'); set_run(r,10,True,MUTED)
    p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); r=p.add_run(description); set_run(r,10,False,INK,True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8); r=p.add_run(f'Figure {number}: {description}'); set_run(r,9,False,MUTED,True)

def page_break(doc): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def main():
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(0.85); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(0.85); sec.right_margin=Inches(0.85); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Aptos'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Aptos'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos'); normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.12
    for n,size,col,bef,aft in [('Heading 1',16,BLUE,16,7),('Heading 2',13,DARK,11,5),('Heading 3',11.5,DARK,8,3)]:
        s=styles[n]; s.font.name='Aptos Display'; s._element.rPr.rFonts.set(qn('w:ascii'),'Aptos Display'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Aptos Display'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
    # header/footer
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('JIO FAQ CHATBOT  |  PROJECT REPORT'); set_run(r,8.5,True,MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_page_field(footer)

    # Cover
    add_para(doc,'PROJECT REPORT',before=35,after=12,align=WD_ALIGN_PARAGRAPH.CENTER,size=13,bold=True,color=BLUE)
    add_para(doc,'Jio FAQ Chatbot',after=7,align=WD_ALIGN_PARAGRAPH.CENTER,size=29,bold=True,color=DARK)
    add_para(doc,'A Mobile AI Assistant for Jio FAQ Support, Document-Grounded Answers, Voice Interaction, and Live Vision',after=22,align=WD_ALIGN_PARAGRAPH.CENTER,size=14,color=MUTED)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(24); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'); set_run(r,13,False,BLUE)
    meta=[('Developed by','Aman Pausker'),('Project type','AI-powered mobile customer-support chatbot'),('Submitted to','[Organisation / Institution name]'),('Mentor / Guide','[Mentor name]'),('Technology stack','React Native (Expo), FastAPI, LangGraph, Neo4j, Qdrant, Supabase, Ollama'),('Date','July 2026')]
    t=doc.add_table(rows=0, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for k,v in meta:
        cells=t.add_row().cells
        for c in cells: borders(c,'FFFFFF'); set_cell_margins(c,70,80,70,80)
        shade(cells[0],'EAF2FB'); p=cells[0].paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_run(p.add_run(k),10,True,DARK)
        p=cells[1].paragraphs[0]; p.paragraph_format.space_after=Pt(0); set_run(p.add_run(v),10,False,INK)
    add_para(doc,'Note: Replace the bracketed academic details with your actual submission information before printing.',before=25,after=0,align=WD_ALIGN_PARAGRAPH.CENTER,size=9,italic=True,color=MUTED)
    page_break(doc)

    # TOC
    add_heading(doc,'Table of Contents',1)
    toc=['1. Executive Summary','2. Problem Statement and Objectives','3. System Overview','4. Requirements and Technology Stack','5. Architecture and Data Flow','6. Core Implementation','7. Mobile Application Experience','8. Data, Security and Privacy','9. Testing and Validation','10. Results and Limitations','11. Future Enhancements','12. Project Structure','13. Screenshot Insertion Guide','14. Conclusion','Appendix A. Suggested Demonstration Script']
    for x in toc: add_para(doc,x,after=3,size=10.5)
    page_break(doc)

    add_heading(doc,'1. Executive Summary',1)
    add_para(doc,'Jio FAQ Chatbot is a mobile-first artificial intelligence assistant designed to answer Jio-related support queries in a conversational format. The project combines a React Native mobile client with a FastAPI service and a retrieval-augmented generation workflow. Instead of depending only on a language model, the system retrieves relevant Jio FAQ content, reranks candidate passages, and instructs the model to answer from the selected context.')
    add_para(doc,'The application supports typed chat, speech-based interaction, document upload for session-specific assistance, image-based queries, weather/tool responses, user memory, chat history, and a live camera experience. The live experience uses WebRTC for media transport and a WebSocket channel for conversational events, enabling a user to speak while showing visual context to the assistant.')
    add_callout(doc,'Project outcome','A prototype customer-support assistant that demonstrates grounded FAQ retrieval, streamed responses, authenticated user sessions, persistent conversational data, personalized memory, and multimodal interaction in one mobile application.')
    add_heading(doc,'Key Contributions',2)
    add_bullets(doc,['Hybrid retrieval that combines vector similarity and full-text search over a Neo4j FAQ graph.','Cross-encoder reranking to select the most relevant FAQ or uploaded-document passage.','A mobile interface built with Expo/React Native for text, voice, image, and live-camera interaction.','Server-sent event streaming for responsive chat output and WebSocket/WebRTC support for live interaction.','User-scoped document retrieval and persistent profiles, sessions, messages, and memories.'])

    add_heading(doc,'2. Problem Statement and Objectives',1)
    add_heading(doc,'2.1 Problem Statement',2)
    add_para(doc,'Customers often need rapid, accurate answers about plans, services, offers, recharge, connectivity, and account-related support. Conventional keyword search makes it difficult to express a question naturally, while a generic language model may answer confidently without using the correct product information. The project addresses this gap by retrieving relevant knowledge before generating an answer and by falling back to a general assistant when a query is unrelated to Jio FAQs.')
    add_heading(doc,'2.2 Objectives',2)
    add_table(doc,['Objective','How the implementation addresses it'],[
        ('Provide natural-language FAQ help','The workflow normalizes common Jio terms, retrieves candidate FAQ entries, reranks them, and generates a concise grounded response.'),
        ('Support multiple user input modes','The mobile client supports typed text, image attachment, voice interaction, document upload, and a live camera session.'),
        ('Maintain relevance and privacy boundaries','Uploaded-document retrieval is filtered by user and session; protected APIs use bearer-token authentication.'),
        ('Keep interactions responsive','Chat replies stream progressively through server-sent events; live interactions use WebSocket events and WebRTC media.'),
        ('Remember useful preferences','A user-memory store preserves approved facts, while chat sessions retain message history and generated titles.')
    ],[2.05,4.45])

    add_heading(doc,'3. System Overview',1)
    add_para(doc,'The solution is organized into three cooperating layers: a mobile client, an AI application service, and storage/retrieval services. The client sends authenticated requests to the API. The API selects either the FAQ route or a general-assistant route. For FAQ questions, it searches the knowledge graph and any user-uploaded material, reranks the results, and passes the selected context to the model. The final response is streamed back to the client and stored with the session.')
    add_heading(doc,'3.1 High-Level Data Flow',2)
    add_table(doc,['Stage','Flow'],[
        ('1. User input','A user enters text, speaks, attaches an image/PDF, or starts the live camera experience.'),
        ('2. Authentication','The mobile app obtains a Supabase session token and includes it in protected API requests.'),
        ('3. Retrieval and routing','The backend normalizes the question, creates an embedding, searches Neo4j and user-scoped Qdrant content, then reranks candidates.'),
        ('4. Answer generation','A LangGraph workflow routes to the FAQ generator or general generator; Ollama streams model output.'),
        ('5. Delivery and persistence','The API sends streamed tokens/final metadata to the app and saves messages, session data, and optional memory updates.')
    ],[1.45,5.05])
    add_para(doc,'\nArchitecture Diagram:\n┌─────────────────────┐\n│    Expo Mobile App   │──── HTTP / WebSocket ────┐\n│  (Gemma 4 on-device) │                           │\n├─────────────────────┤                           ▼\n│    Web Frontend      │                   ┌──────────────────┐\n│   (React + Vite)     │──── HTTP / WS ───▶│  FastAPI Server   │\n└─────────────────────┘                   │  (REST + WebSockets)│\n                                           └────────┬─────────┘\n                                                    │\n                                            LangGraph astream_events()\n                                                    │\n                                        ┌───────────┼───────────────────┐\n                                        ▼           ▼                   ▼\n                                  ┌──────────┐ ┌──────────┐     ┌──────────┐\n                                  │  Neo4j   │ │  Qdrant  │     │  Ollama  │\n                                  │ (vector  │ │ (user    │     │ (Gemma2  │\n                                  │+fulltext │ │  docs)   │     │  Qwen-V) │\n                                  │  hybrid) │ └──────────┘     └──────────┘\n                                  └──────────┘\n                                        │\n                              ┌─────────┴─────────┐\n                              ▼                     ▼\n                        ┌──────────┐         ┌──────────┐\n                        │Supabase  │         │  SQLite   │\n                        │(Auth+DB) │         │(checkpts) │\n                        └──────────┘         └──────────┘')
    add_figure_placeholder(doc,'1','High-level architecture: Mobile App → FastAPI → LangGraph Retrieval/Generation → Neo4j, Qdrant, Supabase, and Ollama')

    add_heading(doc,'4. Requirements and Technology Stack',1)
    add_table(doc,['Layer','Technologies used','Purpose'],[
        ('Mobile application','React Native, Expo, Expo Router, TypeScript','Cross-platform client, navigation, chat UI, camera/audio permissions.'),
        ('API service','FastAPI, Pydantic, Uvicorn-compatible ASGI stack','Authenticated endpoints, streamed chat, file upload, WebSocket sessions.'),
        ('AI orchestration','LangGraph, LangChain, Ollama / ChatOllama','Stateful workflow, model calls, streaming responses, tool-enabled general chat.'),
        ('Retrieval','Sentence Transformers, CrossEncoder, Neo4j, Qdrant','Embeddings, FAQ graph/full-text search, reranking, document retrieval.'),
        ('Persistence','Supabase, SQLite checkpointing','Authentication, chat history, user memory, sessions, workflow checkpoints.'),
        ('Live media','WebRTC, aiortc, WebSocket, Expo AV/Camera','Bidirectional live audio/video and event-based conversation.')
    ],[1.25,2.25,3.0])
    add_heading(doc,'4.1 Minimum Demonstration Environment',2)
    add_bullets(doc,['A device or emulator capable of running the Expo React Native application.','A running FastAPI backend, configured with the required local or hosted service URLs.','Neo4j prepared with FAQ nodes, vector index, and full-text index.','Qdrant configured for user-uploaded document embeddings.','Supabase credentials for authentication and persistent user data.','Ollama running the configured text and optional vision models.'])

    add_heading(doc,'5. Architecture and Data Flow',1)
    add_heading(doc,'5.1 Conversational Workflow',2)
    add_para(doc,'The backend defines a LangGraph state machine with retrieval, FAQ generation, and general generation nodes. The retrieval node first applies lightweight product-name normalization and fuzzy matching. It then encodes the question using an all-MiniLM-L6-v2 embedding model. A Neo4j query combines vector and full-text candidates and traverses the Topic → Subtopic → FAQ relationship. User-uploaded document chunks from Qdrant are added as candidates, after which a cross-encoder scores query–passage pairs.')
    add_para(doc,'If the best candidate score is below the routing threshold, the workflow uses the general-generation node. Otherwise, the FAQ generator receives the selected context and is instructed to answer only from it. Very high-confidence retrievals can be returned directly, avoiding unnecessary generation.')
    add_table(doc,['Node','Responsibility','Output'],[
        ('retrieve','Normalize query, embed it, search Neo4j/Qdrant, rerank candidates, decide the route.','Context, route, and optional direct answer.'),
        ('generate','Generate a Jio answer constrained to the retrieved context.','Grounded FAQ response.'),
        ('general_generation','Handle non-FAQ conversation and permitted tools such as weather/location.','General conversational response.')
    ],[1.35,3.2,1.95])
    add_heading(doc,'5.2 FAQ Knowledge Graph',2)
    add_para(doc,'FAQ content is represented with a graph-oriented structure in which topics contain subtopics and subtopics contain FAQ nodes. This preserves product grouping while allowing semantic and lexical retrieval. The graph query returns the topic, subtopic, question, and answer text for reranking, then uses the selected answer text as context for the language model.')
    add_figure_placeholder(doc,'2','FAQ retrieval flow: Query normalization → embedding + keyword query → Neo4j/Qdrant candidates → cross-encoder reranking → routing decision')

    add_heading(doc,'6. Core Implementation',1)
    add_heading(doc,'6.1 Streaming Text Chat',2)
    add_para(doc,'The /api/chat endpoint accepts a message, optional session ID, and optional image. For ordinary chat, it compiles the workflow with an asynchronous SQLite checkpoint store, listens to model-stream events, and sends incremental token events to the client using the text/event-stream media type. At the end of a response, it returns a final structured payload containing the reply, UI metadata, and session ID.')
    add_heading(doc,'6.2 User Documents',2)
    add_para(doc,'The application accepts PDF uploads and processes them in the background. A document conversion component exports the file to Markdown, a recursive character splitter creates overlapping chunks, and a sentence-transformer encoder generates vector embeddings. Each Qdrant record stores the chunk text together with document, user, and session identifiers. Search is filtered by user ID and session ID, preventing a session from retrieving another user’s uploaded material.')
    add_heading(doc,'6.3 Personal Memory and Session History',2)
    add_para(doc,'The backend saves user and assistant messages to persistent storage and can generate a human-readable title for a new chat session in the background. A separate user-memory record stores explicit facts. Memory is fetched when relevant and becomes part of the system prompt, while the underlying chat workflow retains short-term state through checkpointing.')
    add_heading(doc,'6.4 Multimodal Inputs',2)
    add_bullets(doc,['Image queries: the API compresses the incoming image, sends it to a configured vision model, and returns the result in the chat stream.','Voice interaction: client audio capture and speech transcription are used to turn spoken questions into conversational input; text-to-speech can render responses.','Live camera: the mobile client creates a WebRTC offer, streams microphone/camera tracks, and exchanges interaction events over WebSocket.','Weather card: a tool-enabled general path can transform weather output into structured UI metadata for a rich client-side card.'])

    add_heading(doc,'7. Mobile Application Experience',1)
    add_para(doc,'The Expo application separates authentication and protected application routes. The standard chat screen manages messages, session history, image attachment, loading states, local-model controls, audio state, and account information. The live screen requests camera and microphone permission, opens a WebSocket for events, establishes a WebRTC peer connection, shows the local camera preview, and displays live user/assistant transcript messages.')
    add_table(doc,['Screen / feature','User-visible purpose'],[
        ('Login / authentication','Creates a protected session before access to user data and chat APIs.'),
        ('Chat screen','Displays conversational messages, a text input, image attachment, history, and contextual cards.'),
        ('Voice mode','Lets the user speak and hear replies without relying solely on typing.'),
        ('Account / memory view','Shows the authenticated account identity and saved memory facts.'),
        ('Live Camera','Combines camera preview, microphone input, animated state feedback, and live transcripts.')
    ],[2.0,4.5])
    add_figure_placeholder(doc,'3','Authenticated chat screen showing a question, streamed response, input controls, and session context')
    add_figure_placeholder(doc,'4','Live Camera screen showing camera preview, voice-state animation, and conversational transcript')

    add_heading(doc,'8. Data, Security and Privacy',1)
    add_para(doc,'The prototype uses bearer-token authentication for protected API routes. The API extracts the token from the Authorization header and validates the current user through Supabase. User data is scoped at the application layer: document chunks are stored with user and session identifiers, and Qdrant searches require both identifiers as filters. Chat sessions, messages, and user memory are persisted through authenticated Supabase clients.')
    add_callout(doc,'Important deployment note','This report describes the implemented prototype. Before production deployment, move all credentials out of source code, restrict CORS origins, apply least-privilege database policies, use secure secret management, validate file types and limits, and conduct an application-security review.')
    add_table(doc,['Area','Implemented approach','Recommended hardening'],[
        ('Identity','Bearer token validated through Supabase.','Refresh-token strategy, revocation handling, rate limits, audit logs.'),
        ('Document retrieval','Qdrant payload filters use user and session identifiers.','Encryption at rest, retention policy, malware scanning, stricter content isolation.'),
        ('API transport','HTTP/WebSocket API design with authenticated access.','HTTPS/WSS only, explicit CORS allowlist, request size limits.'),
        ('Model safety','FAQ prompt directs answers to retrieved context.','Output monitoring, refusal policy, quality evaluation set, moderation where required.')
    ],[1.15,2.55,2.8])

    add_heading(doc,'9. Testing and Validation',1)
    add_para(doc,'Testing should validate both functional behavior and user experience. The following plan is aligned with the implemented features and can be used as the basis for project demonstration and evaluation.')
    add_table(doc,['Test area','Example validation','Expected result'],[
        ('FAQ retrieval','Ask a known Jio product/plan question.','The answer follows retrieved FAQ context and does not invent unsupported facts.'),
        ('General routing','Ask a non-Jio conversational question.','The request uses general generation rather than forcing an FAQ answer.'),
        ('Fuzzy normalization','Enter variations such as “jioplus” or a minor product spelling error.','The normalizer maps common variants toward recognized product terms.'),
        ('Document upload','Upload a PDF, then ask a question about its content.','The relevant session-scoped document chunk is retrieved and used as context.'),
        ('Authentication','Call a protected endpoint without or with an invalid token.','The API rejects the request with an authentication error.'),
        ('Streaming','Submit a normal chat query.','Partial output appears before the final event completes.'),
        ('Live interaction','Open the live screen, allow permissions, and speak.','Camera/microphone initialize and transcript/assistant events are shown.'),
        ('History and memory','Start a session and save an approved fact.','Session messages and memory facts remain available to the authenticated user.')
    ],[1.25,2.7,2.55])
    add_heading(doc,'9.1 Evaluation Measures',2)
    add_bullets(doc,['Grounded-answer rate: percentage of FAQ responses supported by the selected context.','Retrieval relevance: manual score of the top retrieved answer for a prepared test-question set.','Response time: time to first streamed token and total completion time.','Task success: percentage of users who obtain a useful answer without reformulating their question.','Multimodal reliability: successful completion rate for image, document, voice, and live-camera flows.'])

    add_heading(doc,'10. Results and Limitations',1)
    add_heading(doc,'10.1 Achieved Capabilities',2)
    add_bullets(doc,['A working architecture for Jio-focused FAQ assistance with semantic retrieval and reranking.','A mobile client that unifies chat, media input, session history, personalized memory, and live interaction.','A server API that streams responses and stores conversational state.','User- and session-scoped retrieval for uploaded documents.','A foundation for structured UI output, demonstrated through a weather-card transformation path.'])
    add_heading(doc,'10.2 Limitations',2)
    add_bullets(doc,['Answer quality depends on the freshness and coverage of the FAQ knowledge base.','Local-model performance and latency depend on available hardware and model configuration.','The live vision flow requires reliable camera/microphone permission, network connectivity, and media interoperability.','Current source configuration should be hardened before public deployment, particularly around secrets and deployment settings.','Formal benchmark results require a controlled test set and repeatable measurement environment.'])
    
    add_heading(doc,'10.3 Performance and Latency Metrics',2)
    add_table(doc,['Metric','Target/Observed Value'],[
        ('Local LLM Inference (General)', '< 2 seconds per response for general agent tasks.'),
        ('Local LLM Inference (Jio FAQ)', '< 10 seconds per response for heavy RAG contexts.'),
        ('Realtime Audio-Chat', '< 2 seconds end-to-end latency for voice responses.'),
        ('Realtime Video-Chat', '< 4 seconds end-to-end latency for vision processing.')
    ],[2.0, 4.0])
    add_heading(doc,'Optimizations for Inference Latency',3)
    add_bullets(doc,[
        'Background Processing: Heavy graph updates (like evaluating and saving Neo4j memories) are deferred to background threads, completely unblocking the main response loop.',
        '"One-Call" Interception: If the LLM triggers a tool call (e.g., fetching weather), the system intercepts the JSON and immediately returns a UI card to the user, skipping a redundant second LLM reasoning pass.',
        'Semantic Routing: By intelligently routing general questions away from the FAQ pipeline, the system avoids unnecessary Qdrant vector lookups and embedding generation.'
    ])

    add_heading(doc,'11. Future Enhancements',1)
    add_table(doc,['Enhancement','Value'],[
        ('Automated FAQ ingestion and change tracking','Keeps the knowledge graph synchronized with approved Jio support content.'),
        ('Multilingual support','Enables Hindi and other regional-language customer interactions.'),
        ('Retrieval evaluation dashboard','Tracks precision, answer grounding, no-answer behavior, and latency across test queries.'),
        ('Human-agent handoff','Escalates low-confidence or sensitive cases with conversational context attached.'),
        ('Production observability','Adds structured metrics, tracing, rate limiting, and alerts for API/retrieval/model health.'),
        ('Richer vision assistance','Uses contextual image analysis for device setup, signal issues, and self-service troubleshooting.')
    ],[2.35,4.15])

    add_heading(doc,'12. Project Structure',1)
    add_para(doc,'The repository is organized around a Python AI backend and an Expo mobile application. The following is a concise logical view of the key project areas.')
    add_table(doc,['Location','Role'],[
        ('server.py','FastAPI server, protected API endpoints, streaming chat, uploads, and live-session handling.'),
        ('app.py / nodes.py','LangGraph workflow and retrieval/generation nodes.'),
        ('system_instructions.py','Prompt builders for general chat, FAQ chat, live voice, and vision modes.'),
        ('file_workflow.py','PDF conversion, chunking, embedding creation, Qdrant storage and search.'),
        ('tools.py','Callable utility tools used by general conversation paths.'),
        ('mobile_app/src/app/','Expo Router screens for authentication, chat, and live camera.'),
        ('mobile_app/src/services/','On-device LiteRT LLM execution, API integrations, and prompts.'),
        ('mobile_app/src/utils/','Supabase client initialization.'),
        ('public/','Static assets served by the backend.')
    ],[2.1,4.4])

    add_heading(doc,'13. Screenshot Insertion Guide',1)
    add_para(doc,'Capture the screenshots below from your running application and replace the matching placeholders in this report. Keep personal phone numbers, access tokens, email addresses, and private document contents hidden or blurred. Use a consistent device frame or clean crop, and place a short caption below each image.')
    add_table(doc,['No.','Screenshot to capture','What it should prove','Suggested section'],[
        ('1','Login / sign-in screen','The application has an authenticated entry point.','Mobile Application Experience'),
        ('2','Main chat screen before sending a query','Chat layout, navigation, and input controls.','Mobile Application Experience'),
        ('3','FAQ query with response','A real Jio-related question and the assistant’s grounded answer.','Executive Summary / Results'),
        ('4','Streaming response in progress','Partial assistant output or loading state before completion.','Core Implementation'),
        ('5','Chat history or new-session view','Persistence and session organization.','Core Implementation'),
        ('6','Account / memory modal','Authenticated profile and saved-memory feature.','Mobile Application Experience'),
        ('7','PDF/document attachment flow','Document selection or upload acknowledgement.','Core Implementation'),
        ('8','Question answered from uploaded document','Session-specific document-grounded assistance.','Core Implementation'),
        ('9','Image attachment and answer','Vision/multimodal chat capability.','Core Implementation'),
        ('10','Voice mode during recording or playback','Speech-based interaction.','Mobile Application Experience'),
        ('11','Live Camera screen','Camera preview, wave animation, and transcript cards.','Mobile Application Experience'),
        ('12','Backend API/terminal log or API documentation','Successful server startup and a streamed request; redact secrets.','Testing and Validation'),
        ('13','Neo4j Browser graph or FAQ node query','Topic → Subtopic → FAQ organization.','Architecture and Data Flow'),
        ('14','Qdrant collection/query result','User-document chunks returned with scoped metadata; blur private text.','Architecture and Data Flow'),
        ('15','Supabase table or dashboard view','Chat sessions/messages/memory persistence; anonymize all user data.','Data, Security and Privacy')
    ],[0.38,1.75,2.7,1.67])
    add_callout(doc,'Best screenshot set for a compact report','If your college expects only 6–8 images, prioritize: login, chat home, Jio FAQ answer, document upload + document answer, voice mode, live camera, Neo4j FAQ graph, and Qdrant/Supabase persistence view.')

    add_heading(doc,'14. Conclusion',1)
    add_para(doc,'Jio FAQ Chatbot demonstrates how a customer-support experience can combine a modern mobile interface with retrieval-grounded AI. Its architecture separates application concerns cleanly: the React Native client manages the interaction, FastAPI provides authenticated streaming services, LangGraph controls the retrieval/generation workflow, and specialized stores support knowledge, documents, sessions, and user memory. The result is a practical prototype that can answer FAQ-style questions while extending naturally to voice, images, uploaded documents, and live visual interaction.')
    add_para(doc,'The next stage is to strengthen operational readiness: curate and evaluate the knowledge base, establish measurable quality targets, and apply production security and observability controls. With these additions, the prototype can evolve into a robust, user-centered support platform.')

    page_break(doc)
    add_heading(doc,'Appendix A. Suggested Demonstration Script',1)
    add_numbered(doc,['Launch the backend and mobile application, then sign in with a test account.','Show the main chat screen and ask a clear Jio FAQ question such as a plan or service query.','Explain that the backend retrieves from the FAQ graph and reranks candidates before answering.','Attach a non-sensitive PDF, wait for processing, and ask a question contained in that document.','Open the account/memory view and show that user-specific information is stored separately.','Demonstrate voice mode or submit an image-based question.','Open Live Camera, grant microphone/camera permission, and show the live transcript response.','Conclude by showing the Neo4j graph or Qdrant collection and explain user/session filtering.'])
    add_heading(doc,'Appendix B. Submission Checklist',1)
    add_bullets(doc,['Replace bracketed cover-page values with your institution, mentor, and submission date.','Insert screenshots from the guide and update each Figure caption.','Remove or blur API keys, tokens, email addresses, phone numbers, private documents, and user IDs.','Verify all features shown in screenshots run on your final demonstration environment.','Update any technology/model names if your final deployed configuration differs.'])
    doc.core_properties.title='Jio FAQ Chatbot Project Report'; doc.core_properties.author='Aman Pausker'; doc.core_properties.subject='Technical project report'
    doc.save(OUT)

if __name__=='__main__': main()
